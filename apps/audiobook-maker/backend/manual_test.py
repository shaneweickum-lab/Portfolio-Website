"""One-off local verification script (not part of the app). Exercises parsing,
job orchestration, and mp3 encoding using a fake voice, since the real Piper
voice model can't be downloaded from this sandbox (huggingface.co is blocked
by the network proxy here). Run with the real voice model once deployed.
"""

import shutil
import sys
import tempfile
import time
from pathlib import Path

import numpy as np

TMP_DATA_DIR = Path(tempfile.mkdtemp(prefix="audiobook-test-"))
import os

os.environ["DATA_DIR"] = str(TMP_DATA_DIR)

sys.path.insert(0, str(Path(__file__).parent))

from app.parsing.docx_doc import parse_docx  # noqa: E402
from app.parsing.markdown_doc import parse_markdown  # noqa: E402
from app.parsing.plain_text import parse_txt  # noqa: E402
import app.tts.voice as voice_module  # noqa: E402


class FakeConfig:
    sample_rate = 22050


class FakeChunk:
    def __init__(self, num_samples: int):
        tone = (np.sin(np.linspace(0, 40, num_samples)) * 3000).astype(np.int16)
        self._bytes = tone.tobytes()

    @property
    def audio_int16_bytes(self) -> bytes:
        return self._bytes


class FakeVoice:
    config = FakeConfig()

    def synthesize(self, text: str):
        yield FakeChunk(num_samples=max(len(text) * 40, 400))


def check(label: str, condition: bool, detail: str = ""):
    status = "PASS" if condition else "FAIL"
    print(f"[{status}] {label}" + (f" -- {detail}" if detail and not condition else ""))
    if not condition:
        raise SystemExit(1)


# --- 1. Parsing ---------------------------------------------------------

txt_path = TMP_DATA_DIR / "sample.txt"
txt_path.write_text("First paragraph.\n\nSecond paragraph here.\n\nThird.", encoding="utf-8")
txt_doc = parse_txt(txt_path)
check("parse_txt splits paragraphs", txt_doc.paragraphs == [
    "First paragraph.", "Second paragraph here.", "Third."
], str(txt_doc.paragraphs))

md_path = TMP_DATA_DIR / "sample.md"
md_path.write_text(
    "# Title\n\nSome **bold** text in a paragraph.\n\n- item one\n- item two\n\n"
    "```\ncode.that.should.be.skipped()\n```\n",
    encoding="utf-8",
)
md_doc = parse_markdown(md_path)
check(
    "parse_markdown extracts text, strips code blocks",
    md_doc.paragraphs == [
        "Title",
        "Some bold text in a paragraph.",
        "item one",
        "item two",
    ],
    str(md_doc.paragraphs),
)

try:
    import docx as docx_lib

    docx_path = TMP_DATA_DIR / "sample.docx"
    document = docx_lib.Document()
    document.add_paragraph("Chapter one begins here.")
    document.add_paragraph("A second paragraph follows.")
    document.save(str(docx_path))

    docx_doc = parse_docx(docx_path)
    check(
        "parse_docx extracts paragraph text",
        docx_doc.paragraphs == ["Chapter one begins here.", "A second paragraph follows."],
        str(docx_doc.paragraphs),
    )
except Exception as exc:  # noqa: BLE001
    check("parse_docx extracts paragraph text", False, str(exc))

# --- 2. Synthesis + mp3 encoding (fake voice) ---------------------------

voice_module._voice = FakeVoice()  # inject fake voice via the lazy-load seam

from app.tts.synthesize import synthesize_to_wav  # noqa: E402

progress_events: list[float] = []
wav_path = TMP_DATA_DIR / "out.wav"
synthesize_to_wav(
    ["Paragraph one.", "Paragraph two.", "Paragraph three."],
    wav_path,
    on_progress=progress_events.append,
)
check("synthesize_to_wav writes a non-empty wav", wav_path.stat().st_size > 44)
check("progress callback fires per paragraph", progress_events == [1 / 3, 2 / 3, 1.0], str(progress_events))

import asyncio  # noqa: E402
from app.tts.synthesize import encode_to_mp3  # noqa: E402

mp3_path = TMP_DATA_DIR / "out.mp3"
asyncio.run(encode_to_mp3(wav_path, mp3_path))
check("encode_to_mp3 produces a non-empty mp3", mp3_path.exists() and mp3_path.stat().st_size > 0)

# --- 3. Full API flow (upload -> poll -> download) ----------------------

from fastapi.testclient import TestClient  # noqa: E402
from app.main import app  # noqa: E402

with TestClient(app) as client:
    with open(txt_path, "rb") as fh:
        response = client.post("/jobs", files={"file": ("sample.txt", fh, "text/plain")})
    check("POST /jobs returns 200", response.status_code == 200, response.text)
    job = response.json()
    check("job starts queued/parsing", job["status"] in ("queued", "parsing"))

    job_id = job["id"]
    final = None
    for _ in range(50):
        status_response = client.get(f"/jobs/{job_id}")
        final = status_response.json()
        if final["status"] in ("done", "error"):
            break
        time.sleep(0.1)

    check("job reaches 'done'", final is not None and final["status"] == "done", str(final))

    download = client.get(f"/jobs/{job_id}/download")
    check("download returns mp3 bytes", download.status_code == 200 and len(download.content) > 0)
    check("download content-type is audio/mpeg", download.headers["content-type"] == "audio/mpeg")

print("\nAll checks passed.")
shutil.rmtree(TMP_DATA_DIR, ignore_errors=True)

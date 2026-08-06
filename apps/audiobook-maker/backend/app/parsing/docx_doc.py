from pathlib import Path

import docx

from app.parsing.base import ParsedDocument

HEADING_STYLE_PREFIX = "Heading"
SKIPPED_STYLES = {"Caption", "Footer", "Header"}


def parse_docx(path: Path) -> ParsedDocument:
    document = docx.Document(str(path))

    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name if paragraph.style else "") or ""
        if style_name in SKIPPED_STYLES:
            continue

        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    return ParsedDocument(title=path.stem, paragraphs=paragraphs)
